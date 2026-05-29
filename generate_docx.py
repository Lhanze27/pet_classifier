from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Default body style ────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_heading(para, level, text, color=None):
    para.style = doc.styles[f'Heading {level}']
    run = para.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def add_heading(text, level, color=None):
    p = doc.add_paragraph()
    set_heading(p, level, text, color)
    return p

def add_paragraph(text='', bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
    return p

def add_code_block(code_text):
    """Light-grey shaded paragraph for code."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    # grey background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    # handle inline **bold** or `code`
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2]); run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1]); run.font.name = 'Courier New'
        else:
            p.add_run(part)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        # header bg color (dark blue)
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '2E4057')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # data rows
    for r, row in enumerate(rows):
        tr = table.rows[r + 1]
        for c, cell_text in enumerate(row):
            cell = tr.cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(cell_text)
            run.font.size = Pt(10)
            if (r % 2) == 1:
                tc   = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd  = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  'EAF0FB')
                tcPr.append(shd)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    doc.add_paragraph()   # spacing after table
    return table

# ═══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('Pet Breed Classifier')
title_run.font.size  = Pt(28)
title_run.font.bold  = True
title_run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run('Detailed User Manual')
sub_run.font.size  = Pt(18)
sub_run.font.color.rgb = RGBColor(0x55, 0x6B, 0x8D)

doc.add_paragraph()
meta_lines = [
    ('System:',   'AI-powered Pet Breed Classification using Deep Learning & Transfer Learning'),
    ('Model:',    'EfficientNetV2B0 (37 breed classes)'),
    ('Accuracy:', '~90–95%'),
    ('Est. Time:','~1.5 hours (mostly training)'),
]
for label, value in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + '  '); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(value);        r2.font.size = Pt(11)

doc.add_paragraph()

# ── Horizontal rule ───────────────────────────────────────────────────────────
hr = doc.add_paragraph()
hr.paragraph_format.space_after = Pt(2)
pPr = hr._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'),   'single')
bottom.set(qn('w:sz'),    '6')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), '2E4057')
pBdr.append(bottom)
pPr.append(pBdr)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (plain text listing)
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('Table of Contents', 1, color=(0x2E, 0x40, 0x57))
toc_items = [
    '1.  Prerequisites',
    '2.  Getting the Code from GitHub',
    '3.  Python Installation',
    '4.  GPU Setup (NVIDIA CUDA & cuDNN)',
    '5.  Setting Up the Virtual Environment',
    '6.  Installing Dependencies',
    '7.  Verifying Your Setup',
    '8.  Training the Model',
    '9.  Evaluating the Model & Generating Figures',
    '10. Understanding the Output Files',
    '11. Troubleshooting',
    '12. Maintenance Tips',
    '13. Quick Run Cheat Sheet',
    '14. Contact & Support',
]
for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.add_run(item).font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — PREREQUISITES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('1. Prerequisites', 1, color=(0x2E, 0x40, 0x57))
add_paragraph('Before starting, make sure your system meets the following requirements.')

add_heading('Minimum Requirements', 2)
add_table(
    ['Component', 'Requirement'],
    [
        ['Operating System', 'Windows 10 or Windows 11'],
        ['Python Version',   '3.10 or 3.11  (do NOT use 3.12+)'],
        ['RAM',              '8 GB'],
        ['GPU',              'NVIDIA GPU (strongly recommended)'],
        ['Storage',          'At least 5 GB of free disk space'],
    ],
    col_widths=[2.0, 3.8]
)

add_heading('Recommended Requirements', 2)
add_table(
    ['Component', 'Recommendation'],
    [
        ['GPU',          'NVIDIA RTX 4050 (6 GB VRAM)'],
        ['CUDA Version', '11.8'],
        ['RAM',          '16 GB'],
        ['Internet',     'Stable connection (~800 MB dataset download)'],
    ],
    col_widths=[2.0, 3.8]
)

add_heading('Required Software (installed during this guide)', 2)
for item in [
    'Python 3.10 or 3.11',
    'Git (for cloning/forking from GitHub)',
    'NVIDIA GPU Drivers',
    'CUDA Toolkit 11.8',
    'cuDNN 8.6',
]:
    add_bullet(item)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — GETTING THE CODE FROM GITHUB
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('2. Getting the Code from GitHub', 1, color=(0x2E, 0x40, 0x57))
add_paragraph(
    'You have two options: Cloning (direct copy, no GitHub account needed) or '
    'Forking (creates your own copy on GitHub, requires an account). Choose one.'
)

add_heading('Option A — Clone the Repository (Recommended for most users)', 2)
add_paragraph(
    'Cloning gives you a local copy of the project so you can run it on your machine.'
)

add_heading('Step 1: Install Git', 3)
for line in [
    'Go to https://git-scm.com/downloads',
    'Download the Windows installer and run it.',
    'During installation, accept all defaults and click Next through each screen.',
    'Verify Git is installed by opening Command Prompt and running:',
]:
    add_bullet(line)
add_code_block('git --version')
add_bullet('You should see something like: git version 2.x.x')

add_heading('Step 2: Navigate to where you want the project', 3)
add_paragraph('Open Command Prompt (press Windows key, type cmd, press Enter):')
add_code_block('cd Desktop')

add_heading('Step 3: Clone the repository', 3)
add_code_block('git clone https://github.com/lhanze27/pet_classifier.git')

add_heading('Step 4: Move into the project folder', 3)
add_code_block('cd pet_classifier')
add_paragraph('Your project folder is now ready at Desktop\\pet_classifier.')

add_heading('Option B — Fork the Repository (Recommended for contributors)', 2)
add_paragraph(
    'Forking creates your own copy of the repository on your GitHub account, '
    'allowing you to make changes and push them independently.'
)

add_heading('Step 1: Create a GitHub account (if you don\'t have one)', 3)
for line in [
    'Go to https://github.com',
    'Click Sign up and follow the on-screen steps.',
]:
    add_bullet(line)

add_heading('Step 2: Fork the repository', 3)
for line in [
    'Go to the original repository page on GitHub.',
    'Click the Fork button in the top-right corner of the page.',
    'Under "Owner", select your GitHub username.',
    'Click Create fork.',
]:
    add_bullet(line)

add_heading('Step 3: Install Git', 3)
add_paragraph('Follow the same steps as Option A, Step 1.')

add_heading('Step 4: Clone your fork to your local machine', 3)
add_paragraph('Open Command Prompt:')
add_code_block('cd Desktop\ngit clone https://github.com/YOUR-USERNAME/pet_classifier.git\ncd pet_classifier')
add_paragraph('Replace YOUR-USERNAME with your actual GitHub username.')

add_heading('Step 5: (Optional) Connect to the original repo for future updates', 3)
add_code_block('git remote add upstream https://github.com/lhanze27/pet_classifier.git')
add_paragraph('This lets you pull future updates from the original project using:')
add_code_block('git fetch upstream\ngit merge upstream/main')

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — PYTHON INSTALLATION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('3. Python Installation', 1, color=(0x2E, 0x40, 0x57))
p = doc.add_paragraph()
r = p.add_run('IMPORTANT: ')
r.bold = True; r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
p.add_run('Only use Python 3.10 or 3.11. Python 3.12+ is NOT compatible with TensorFlow 2.10.')

add_heading('Step 1: Download Python', 2)
for line in [
    'Go to https://www.python.org/downloads/',
    'Under "Looking for a specific release?", scroll down and download Python 3.11.x or Python 3.10.x.',
    'Click the Windows installer link (e.g., Windows installer (64-bit)).',
]:
    add_bullet(line)

add_heading('Step 2: Run the installer', 2)
for line in [
    'Double-click the downloaded .exe file.',
    'On the first screen, CHECK the box that says "Add Python to PATH" — this is critical.',
    'Click Install Now.',
    'Wait for the installation to complete and click Close.',
]:
    add_bullet(line)

add_heading('Step 3: Verify Python is installed correctly', 2)
add_paragraph('Open a new Command Prompt window and run:')
add_code_block('python --version')
add_paragraph('You should see:  Python 3.11.x  or  Python 3.10.x')
add_paragraph('If you see Python 3.12.x or higher, uninstall and reinstall the correct version.')

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — GPU SETUP
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('4. GPU Setup (NVIDIA CUDA & cuDNN)', 1, color=(0x2E, 0x40, 0x57))
add_paragraph(
    'This section configures your NVIDIA GPU to work with TensorFlow for accelerated training. '
    'Skip this section only if you plan to train on CPU (which will be significantly slower).'
)

add_heading('Step 1: Update NVIDIA GPU Drivers', 2)
for line in [
    'Download GeForce Experience from nvidia.com/en-us/geforce/geforce-experience/ and install it.',
    'Open GeForce Experience.',
    'Go to the Drivers tab and click Check for updates.',
    'Download and install the latest driver.',
    'Restart your computer after the driver installs.',
]:
    add_bullet(line)

add_heading('Step 2: Install CUDA Toolkit 11.8', 2)
for line in [
    'Go to the CUDA 11.8 archive: developer.nvidia.com/cuda-11-8-0-download-archive',
    'Select: Operating System → Windows | Architecture → x86_64 | Version → 11 | Installer Type → exe (local)',
    'Click Download (the file is approximately 2.5 GB).',
    'Run the downloaded installer.',
    'Choose Express (Recommended) installation and click Next.',
    'Wait for the installation to complete (~10–15 minutes).',
    'Restart your computer.',
]:
    add_bullet(line)

add_heading('Step 3: Install cuDNN 8.6 for CUDA 11.8', 2)
for line in [
    'Go to the cuDNN archive: developer.nvidia.com/rdp/cudnn-archive',
    'You will need a free NVIDIA Developer Account to download.',
    'Find "cuDNN v8.6.x for CUDA 11.x" and download the Windows zip file.',
    'Extract the downloaded zip file.',
    'Inside the extracted folder, you will see three folders: bin, include, and lib.',
    'Copy the contents of each folder into the corresponding folder inside:',
]:
    add_bullet(line)
add_code_block('C:\\Program Files\\NVIDIA GPU Computing Toolkit\\CUDA\\v11.8\\')
for line in [
    'Copy contents of bin    →  ...\\CUDA\\v11.8\\bin\\',
    'Copy contents of include →  ...\\CUDA\\v11.8\\include\\',
    'Copy contents of lib    →  ...\\CUDA\\v11.8\\lib\\',
    'Restart your computer.',
]:
    add_bullet(line)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — VIRTUAL ENVIRONMENT
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('5. Setting Up the Virtual Environment', 1, color=(0x2E, 0x40, 0x57))
add_paragraph(
    'A virtual environment isolates this project\'s dependencies from other Python projects on your machine.'
)

add_heading('Step 1: Open Command Prompt and navigate to the project folder', 2)
add_code_block('cd Desktop\\pet_classifier')

add_heading('Step 2: Create the virtual environment', 2)
add_code_block('python -m venv venv')
add_paragraph('This creates a venv folder inside your project directory.')

add_heading('Step 3: Activate the virtual environment', 2)
add_code_block('venv\\Scripts\\activate')
add_paragraph(
    'You will know it is active when you see (venv) at the beginning of your command prompt line, like this:'
)
add_code_block('(venv) C:\\Users\\YourName\\Desktop\\pet_classifier>')

p = doc.add_paragraph()
r = p.add_run('Important: ')
r.bold = True
p.add_run(
    'You must activate the virtual environment every time you open a new '
    'Command Prompt window before running any project scripts.'
)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — INSTALLING DEPENDENCIES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('6. Installing Dependencies', 1, color=(0x2E, 0x40, 0x57))
add_paragraph('With the virtual environment active, install all required Python libraries.')

add_heading('Step 1: Upgrade pip to the latest version', 2)
add_code_block('pip install --upgrade pip')

add_heading('Step 2: Install all project dependencies', 2)
add_code_block('pip install -r requirements.txt')

add_paragraph('This will install the following libraries:')
add_table(
    ['Library', 'Version', 'Purpose'],
    [
        ['tensorflow',          '2.10.0',    'Deep learning framework'],
        ['tensorflow-datasets', '4.8.3',     'Downloads and manages the pet dataset'],
        ['matplotlib',          '≥ 3.7.0',   'Generates training curve graphs'],
        ['seaborn',             '≥ 0.12.0',  'Generates the confusion matrix'],
        ['scikit-learn',        '≥ 1.3.0',   'Classification metrics (precision, recall, F1)'],
        ['numpy',               '≥ 1.24.0',  'Numerical operations'],
        ['Pillow',              '≥ 10.0.0',  'Image loading and processing'],
    ],
    col_widths=[1.8, 1.2, 2.8]
)
add_paragraph('Note: This download is approximately 2 GB and may take 5–10 minutes depending on your internet speed.')

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — VERIFYING SETUP
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('7. Verifying Your Setup', 1, color=(0x2E, 0x40, 0x57))
add_paragraph('Before training, verify that Python and the GPU are configured correctly.')

add_heading('Step 1: Verify TensorFlow can see your GPU', 2)
add_paragraph('With the virtual environment active, run:')
add_code_block('python -c "import tensorflow as tf; print(\'GPUs:\', tf.config.list_physical_devices(\'GPU\'))"')

add_paragraph('Expected output (GPU working correctly):')
add_code_block("GPUs: [PhysicalDevice(name='/physical_device:GPU:0', device_type='GPU')]")

add_paragraph('Output if GPU is not detected:')
add_code_block('GPUs: []')

add_paragraph(
    'If the GPU is not detected, revisit Section 4 and ensure CUDA and cuDNN are installed correctly.'
)

add_heading('Step 2: Verify all required files are present', 2)
add_paragraph('Your project folder should contain:')
add_code_block(
    'pet_classifier/\n'
    '├── 01_train_model.py\n'
    '├── 02_evaluate_model.py\n'
    '├── requirements.txt\n'
    '├── README.md\n'
    '└── venv/'
)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — TRAINING THE MODEL
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('8. Training the Model', 1, color=(0x2E, 0x40, 0x57))

add_paragraph('Before you begin training:')
for line in [
    'Plug your laptop into a power outlet (training uses significant power)',
    'Use a cooling pad if available (the GPU will run hot)',
    'Close any heavy applications (games, video editors, browsers with many tabs)',
    'Ensure you have a stable internet connection for the initial dataset download',
]:
    add_bullet(line)

add_heading('Step 1: Make sure your virtual environment is active', 2)
add_code_block('cd Desktop\\pet_classifier\nvenv\\Scripts\\activate')
add_paragraph('You should see (venv) at the start of your prompt.')

add_heading('Step 2: Start the training script', 2)
add_code_block('python 01_train_model.py')

add_heading('Step 3: What to expect during training', 2)
add_paragraph('The script runs in three stages:')
add_table(
    ['Stage', 'Description', 'Estimated Time'],
    [
        ['Dataset Download', 'Downloads the Oxford-IIIT Pet Dataset (~800 MB)', '5–10 min (first run only)'],
        ['Phase 1 Training', 'Trains only the classifier head (base model frozen) for 10 epochs', '~15 minutes'],
        ['Phase 2 Training', 'Fine-tunes the entire model for additional epochs', '~20–30 minutes'],
    ],
    col_widths=[1.5, 3.0, 1.8]
)
for line in [
    'During Phase 1, accuracy should climb from ~3% up to ~85–90%.',
    'During Phase 2, accuracy should improve to ~92–95%.',
    'Training is complete when you see the message: TRAINING COMPLETE!',
]:
    add_bullet(line)

add_heading('Step 4: Do not close the terminal', 2)
add_paragraph(
    'Do not close the Command Prompt window while training is running. '
    'You may minimize it, but closing it will terminate the training process.'
)

add_heading('If training crashes with an Out of Memory error:', 2)
for line in [
    'Open 01_train_model.py in any text editor (Notepad, VS Code, etc.).',
    'Find the line: BATCH_SIZE = 16',
    'Change it to: BATCH_SIZE = 8',
    'Save the file and re-run python 01_train_model.py',
]:
    add_bullet(line)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — EVALUATING THE MODEL
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('9. Evaluating the Model & Generating Figures', 1, color=(0x2E, 0x40, 0x57))
add_paragraph('After training completes successfully, run the evaluation script.')

add_heading('Step 1: Run the evaluation script', 2)
add_paragraph('With the virtual environment still active:')
add_code_block('python 02_evaluate_model.py')

add_heading('Step 2: What the evaluation script produces', 2)
add_table(
    ['Output File', 'Location', 'Purpose'],
    [
        ['training_curves.png',    'outputs/figures/', 'Training and validation accuracy per epoch'],
        ['confusion_matrix.png',   'outputs/figures/', 'Prediction performance across all 37 breeds'],
        ['sample_predictions.png', 'outputs/figures/', 'Predicted vs. actual breed label examples'],
        ['per_class_accuracy.png', 'outputs/figures/', 'Bar chart of accuracy per breed class'],
        ['classification_report.txt', 'outputs/',     'Full precision, recall, and F1-score table'],
        ['results_summary.txt',    'outputs/',         'High-level summary of overall accuracy'],
    ],
    col_widths=[2.0, 1.5, 2.8]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — UNDERSTANDING OUTPUT FILES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('10. Understanding the Output Files', 1, color=(0x2E, 0x40, 0x57))

add_heading('Training Curves (training_curves.png)', 2)
add_paragraph(
    'A line graph showing how training accuracy and validation accuracy changed across each epoch. '
    'Ideally, both lines should increase and converge close together, indicating the model is learning without overfitting.'
)

add_heading('Confusion Matrix (confusion_matrix.png)', 2)
add_paragraph(
    'A grid where rows represent the true breed labels and columns represent the predicted labels. '
    'A perfect classifier would show a bright diagonal from top-left to bottom-right. '
    'Off-diagonal cells indicate misclassifications between visually similar breeds.'
)

add_heading('Per-Class Accuracy (per_class_accuracy.png)', 2)
add_paragraph(
    'A horizontal bar chart showing how accurately the model identifies each of the 37 breeds individually. '
    'Breeds with shorter bars may benefit from additional training data.'
)

add_heading('Sample Predictions (sample_predictions.png)', 2)
add_paragraph(
    'A grid of example images showing the pet photo, the predicted breed, and the actual breed. '
    'Green labels indicate correct predictions; red labels indicate incorrect ones.'
)

add_heading('Classification Report (classification_report.txt)', 2)
add_paragraph('A detailed text table showing three metrics for each breed class:')
for line in [
    'Precision: Of all times the model predicted this breed, what percentage was correct.',
    'Recall: Of all actual images of this breed, what percentage the model correctly identified.',
    'F1-score: The harmonic mean of precision and recall (higher is better; max is 1.0).',
]:
    add_bullet(line)

add_heading('Final Output Folder Structure', 2)
add_code_block(
    'pet_classifier/\n'
    '├── 01_train_model.py\n'
    '├── 02_evaluate_model.py\n'
    '├── requirements.txt\n'
    '├── README.md\n'
    '├── venv/\n'
    '└── outputs/\n'
    '    ├── models/\n'
    '    │   ├── best_phase1.keras\n'
    '    │   ├── best_final.keras\n'
    '    │   └── final_model.keras   ← Main trained model file\n'
    '    ├── figures/\n'
    '    │   ├── training_curves.png\n'
    '    │   ├── confusion_matrix.png\n'
    '    │   ├── sample_predictions.png\n'
    '    │   └── per_class_accuracy.png\n'
    '    ├── class_names.txt\n'
    '    ├── classification_report.txt\n'
    '    └── results_summary.txt'
)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('11. Troubleshooting', 1, color=(0x2E, 0x40, 0x57))

issues = [
    (
        '"No module named \'tensorflow\'"',
        'The virtual environment is not activated.',
        ['Run venv\\Scripts\\activate before running any script.'],
    ),
    (
        '"Could not load dynamic library \'cudart64_110.dll\'"',
        'CUDA is not installed or is installed incorrectly.',
        ['Reinstall CUDA Toolkit 11.8 following Section 4, Step 2 exactly.'],
    ),
    (
        '"ResourceExhaustedError: OOM" (Out of Memory)',
        'The GPU does not have enough VRAM for the current batch size.',
        [
            'Open 01_train_model.py in a text editor.',
            'Locate the line BATCH_SIZE = 16.',
            'Change it to BATCH_SIZE = 8 (or BATCH_SIZE = 4 if the error persists).',
            'Save and rerun python 01_train_model.py',
        ],
    ),
    (
        'GPU not detected (GPUs: [])',
        'NVIDIA drivers, CUDA, or cuDNN are missing or mismatched.',
        [
            'Reinstall the latest NVIDIA drivers via GeForce Experience.',
            'Reinstall CUDA Toolkit 11.8.',
            'Reinstall cuDNN 8.6 and ensure files are copied to the correct CUDA directory.',
            'Restart your computer and retest.',
        ],
    ),
    (
        'Training stuck at low accuracy (below 50%)',
        'The model may be training on CPU instead of GPU.',
        [
            'Open Task Manager (Ctrl + Shift + Esc).',
            'Click the Performance tab and then GPU.',
            'If GPU usage is 0% during training, CUDA is not configured correctly — redo Section 4.',
        ],
    ),
    (
        '"Kernel died" or laptop freezes during training',
        'Excessive VRAM or RAM usage.',
        [
            'Lower BATCH_SIZE in 01_train_model.py (try 8, then 4).',
            'Close all browser tabs and background applications before training.',
        ],
    ),
    (
        'Dataset download fails or hangs',
        'Slow or interrupted internet connection.',
        [
            'Check your internet connection.',
            'Delete the partial download folder: C:\\Users\\YourName\\tensorflow_datasets',
            'Rerun python 01_train_model.py',
        ],
    ),
    (
        'Missing dependencies after pulling updates from GitHub',
        'New packages may have been added to requirements.txt.',
        [
            'With the virtual environment active, run:  pip install --upgrade pip',
            'Then run:  pip install -r requirements.txt',
        ],
    ),
]

for error, cause, fixes in issues:
    p = doc.add_paragraph()
    r = p.add_run(error)
    r.bold = True; r.font.size = Pt(11)

    p2 = doc.add_paragraph()
    r2 = p2.add_run('Cause: ')
    r2.bold = True
    p2.add_run(cause)

    p3 = doc.add_paragraph()
    r3 = p3.add_run('Fix:')
    r3.bold = True

    for fix in fixes:
        add_bullet(fix)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 12 — MAINTENANCE TIPS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('12. Maintenance Tips', 1, color=(0x2E, 0x40, 0x57))
for tip in [
    'Keep TensorFlow updated within the compatible version range (2.10.x) by checking for patch releases.',
    'Back up trained models regularly — copy the outputs/models/ folder to an external drive or cloud storage. The main model file (final_model.keras) is approximately 30 MB.',
    'Store datasets securely — do not delete the tensorflow_datasets folder in your home directory if you plan to retrain, as re-downloading takes 5–10 minutes.',
    'Monitor GPU temperatures during training using MSI Afterburner or GPU-Z. Sustained temperatures above 90°C may indicate insufficient cooling.',
    'Pull the latest updates from GitHub periodically to get bug fixes and improvements.',
]:
    add_bullet(tip)

add_paragraph('To update your local code (clone):')
add_code_block('git pull origin main')
add_paragraph('To update your local code (fork):')
add_code_block('git fetch upstream\ngit merge upstream/main')

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 13 — QUICK RUN CHEAT SHEET
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('13. Quick Run Cheat Sheet', 1, color=(0x2E, 0x40, 0x57))
add_paragraph('Once your one-time setup is complete, running the full pipeline again is as simple as:')
add_code_block(
    'cd Desktop\\pet_classifier\n'
    'venv\\Scripts\\activate\n'
    'python 01_train_model.py\n'
    'python 02_evaluate_model.py'
)

add_paragraph('To update your local code from GitHub before running:')
add_code_block(
    'cd Desktop\\pet_classifier\n'
    'git pull origin main\n'
    'venv\\Scripts\\activate\n'
    'pip install -r requirements.txt\n'
    'python 01_train_model.py\n'
    'python 02_evaluate_model.py'
)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 14 — CONTACT & SUPPORT
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('14. Contact & Support', 1, color=(0x2E, 0x40, 0x57))
add_paragraph('For technical issues or questions, contact the development team:')
add_table(
    ['Name', 'Contact Number'],
    [
        ['Enano, Arcejay',        '0961 122 7283'],
        ['Abunda, Lhanze Tyler',  '0908 653 7095'],
    ],
    col_widths=[2.5, 2.5]
)

add_paragraph('When reporting an issue, please include:')
for item in [
    'A screenshot of the full error message in the Command Prompt',
    'Your Python version (python --version)',
    'Whether your GPU was detected (tf.config.list_physical_devices(\'GPU\'))',
    'Which script produced the error (01_train_model.py or 02_evaluate_model.py)',
]:
    add_bullet(item)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Pet Breed Classifier — User Manual v1.0')
r.italic = True
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ───────────────────────────────────────────────────────────────────────
out = '/home/user/pet_classifier/Pet_Breed_Classifier_User_Manual.docx'
doc.save(out)
print('Saved:', out)
